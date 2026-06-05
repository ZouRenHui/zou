"""Image processing: watermark, resize, edit, OCR."""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Callable, Literal

from PIL import Image, ImageDraw, ImageEnhance, ImageFont, ImageOps

from image_utils import (
    IMAGE_EXTENSIONS,
    find_system_font,
    format_file_size,
    resolve_output_path,
)

Position = Literal[
    "top-left",
    "top-center",
    "top-right",
    "center-left",
    "center",
    "center-right",
    "bottom-left",
    "bottom-center",
    "bottom-right",
    "tile",
]


@dataclass
class WatermarkOptions:
    mode: Literal["text", "logo"] = "text"
    text: str = "水印"
    logo_path: Path | None = None
    font_size: int = 36
    opacity: float = 0.5
    position: Position = "bottom-right"
    margin: int = 20
    color: tuple[int, int, int] = (255, 255, 255)
    logo_scale: float = 0.2


@dataclass
class ResizeOptions:
    width: int | None = None
    height: int | None = None
    keep_aspect: bool = True
    max_side: int | None = None
    quality: int = 92


@dataclass
class CropBox:
    left: int
    top: int
    right: int
    bottom: int


@dataclass
class RemoveWatermarkOptions:
    box: CropBox
    inpaint_radius: int = 5
    method: Literal["telea", "ns"] = "telea"
    mask_expand: int = 2


@dataclass
class ConvertOptions:
    target_ext: str
    quality: int = 85
    jpeg_bg: tuple[int, int, int] = (255, 255, 255)


@dataclass
class CompressOptions:
    quality: int = 85
    max_side: int | None = None
    force_jpeg: bool = False
    jpeg_bg: tuple[int, int, int] = (255, 255, 255)


@dataclass
class CompressResult:
    path: Path
    original_bytes: int
    output_bytes: int

    @property
    def saved_ratio(self) -> float:
        if self.original_bytes <= 0:
            return 0.0
        return (1 - self.output_bytes / self.original_bytes) * 100


def _load_rgba(path: Path) -> Image.Image:
    img = Image.open(path)
    if img.mode not in ("RGBA", "RGB"):
        img = img.convert("RGBA")
    elif img.mode == "RGB":
        img = img.convert("RGBA")
    return img


def save_image(img: Image.Image, dest: Path, quality: int = 92) -> Path:
    """Save a PIL image to disk."""
    return _save_image(img, dest, quality)


def _prepare_for_format(img: Image.Image, ext: str, jpeg_bg: tuple[int, int, int] = (255, 255, 255)) -> Image.Image:
    ext = ext.lower()
    if ext in (".jpg", ".jpeg", ".bmp"):
        if img.mode in ("RGBA", "LA", "P"):
            background = Image.new("RGB", img.size, jpeg_bg)
            if img.mode == "P":
                img = img.convert("RGBA")
            if img.mode in ("RGBA", "LA"):
                background.paste(img, mask=img.split()[-1])
            else:
                background.paste(img)
            return background
        if img.mode != "RGB":
            return img.convert("RGB")
    elif ext == ".gif":
        if img.mode != "P":
            return img.convert("P", palette=Image.Palette.ADAPTIVE)
    elif ext in (".png", ".webp", ".tiff", ".tif"):
        if img.mode not in ("RGB", "RGBA"):
            return img.convert("RGBA" if ext == ".png" else "RGB")
    return img


def _save_image(img: Image.Image, dest: Path, quality: int = 92) -> Path:
    dest.parent.mkdir(parents=True, exist_ok=True)
    ext = dest.suffix.lower()
    prepared = _prepare_for_format(img, ext)

    if ext in (".jpg", ".jpeg"):
        prepared.save(dest, quality=quality, optimize=True, progressive=True)
    elif ext == ".webp":
        prepared.save(dest, quality=quality, method=6)
    elif ext == ".png":
        prepared.save(dest, optimize=True, compress_level=min(9, max(0, (100 - quality) // 11)))
    elif ext in (".tif", ".tiff"):
        prepared.save(dest, compression="tiff_lzw")
    else:
        prepared.save(dest)
    return dest


def _paste_position(
    base_w: int,
    base_h: int,
    overlay_w: int,
    overlay_h: int,
    position: Position,
    margin: int,
) -> tuple[int, int]:
    positions = {
        "top-left": (margin, margin),
        "top-center": ((base_w - overlay_w) // 2, margin),
        "top-right": (base_w - overlay_w - margin, margin),
        "center-left": (margin, (base_h - overlay_h) // 2),
        "center": ((base_w - overlay_w) // 2, (base_h - overlay_h) // 2),
        "center-right": (base_w - overlay_w - margin, (base_h - overlay_h) // 2),
        "bottom-left": (margin, base_h - overlay_h - margin),
        "bottom-center": ((base_w - overlay_w) // 2, base_h - overlay_h - margin),
        "bottom-right": (base_w - overlay_w - margin, base_h - overlay_h - margin),
    }
    return positions.get(position, positions["bottom-right"])


def _apply_opacity(image: Image.Image, opacity: float) -> Image.Image:
    if image.mode != "RGBA":
        image = image.convert("RGBA")
    alpha = image.split()[3]
    alpha = ImageEnhance.Brightness(alpha).enhance(max(0.0, min(1.0, opacity)))
    image.putalpha(alpha)
    return image


def _make_text_overlay(text: str, font_size: int, color: tuple[int, int, int], opacity: float) -> Image.Image:
    font_path = find_system_font()
    try:
        font = ImageFont.truetype(font_path, font_size) if font_path else ImageFont.load_default()
    except OSError:
        font = ImageFont.load_default()

    bbox = font.getbbox(text)
    tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]
    overlay = Image.new("RGBA", (tw + 4, th + 4), (0, 0, 0, 0))
    draw = ImageDraw.Draw(overlay)
    draw.text((2, 2), text, font=font, fill=(*color, int(255 * opacity)))
    return overlay


def add_watermark(source: Path, dest: Path, options: WatermarkOptions) -> Path:
    base = _load_rgba(source)

    if options.mode == "text":
        if not options.text.strip():
            raise ValueError("水印文字不能为空")
        overlay = _make_text_overlay(options.text, options.font_size, options.color, options.opacity)
    else:
        if not options.logo_path or not options.logo_path.is_file():
            raise ValueError("请选择有效的水印 Logo 文件")
        logo = Image.open(options.logo_path).convert("RGBA")
        target_w = max(1, int(base.width * options.logo_scale))
        ratio = target_w / logo.width
        target_h = max(1, int(logo.height * ratio))
        logo = logo.resize((target_w, target_h), Image.Resampling.LANCZOS)
        overlay = _apply_opacity(logo, options.opacity)

    if options.position == "tile":
        tiled = Image.new("RGBA", base.size, (0, 0, 0, 0))
        step_x = max(overlay.width + options.margin, overlay.width)
        step_y = max(overlay.height + options.margin, overlay.height)
        for y in range(0, base.height, step_y):
            for x in range(0, base.width, step_x):
                tiled.paste(overlay, (x, y), overlay)
        overlay = tiled

    result = base.copy()
    if options.position != "tile":
        x, y = _paste_position(
            base.width, base.height, overlay.width, overlay.height, options.position, options.margin
        )
        result.paste(overlay, (x, y), overlay)
    else:
        result = Image.alpha_composite(result, overlay)

    return _save_image(result, dest)


def clamp_crop_box(box: CropBox, img_w: int, img_h: int) -> CropBox:
    left = max(0, min(box.left, img_w - 1))
    top = max(0, min(box.top, img_h - 1))
    right = max(left + 1, min(box.right, img_w))
    bottom = max(top + 1, min(box.bottom, img_h))
    return CropBox(left, top, right, bottom)


def watermark_region_from_preset(
    img_w: int,
    img_h: int,
    position: Position,
    margin: int,
    region_w: int,
    region_h: int,
) -> CropBox:
    if position == "tile":
        raise ValueError("平铺水印请改用「自定义区域」或在预览中框选")
    region_w = max(1, region_w)
    region_h = max(1, region_h)
    x, y = _paste_position(img_w, img_h, region_w, region_h, position, margin)
    return clamp_crop_box(CropBox(x, y, x + region_w, y + region_h), img_w, img_h)


def remove_watermark(source: Path, dest: Path, options: RemoveWatermarkOptions) -> Path:
    try:
        import cv2
        import numpy as np
    except ImportError as exc:
        raise RuntimeError("未安装 opencv-python-headless，请运行: pip install opencv-python-headless") from exc

    img = Image.open(source).convert("RGB")
    w, h = img.size
    box = clamp_crop_box(options.box, w, h)

    arr = np.array(img)
    mask = np.zeros((h, w), dtype=np.uint8)
    mask[box.top : box.bottom, box.left : box.right] = 255

    if options.mask_expand > 0:
        k = options.mask_expand * 2 + 1
        kernel = np.ones((k, k), np.uint8)
        mask = cv2.dilate(mask, kernel, iterations=1)

    flag = cv2.INPAINT_TELEA if options.method == "telea" else cv2.INPAINT_NS
    bgr = cv2.cvtColor(arr, cv2.COLOR_RGB2BGR)
    repaired = cv2.inpaint(bgr, mask, options.inpaint_radius, flag)
    rgb = cv2.cvtColor(repaired, cv2.COLOR_BGR2RGB)
    return _save_image(Image.fromarray(rgb), dest)


def resolve_remove_box(
    source: Path,
    *,
    preset: bool,
    position: Position,
    margin: int,
    region_w: int,
    region_h: int,
    custom_box: CropBox,
) -> CropBox:
    with Image.open(source) as img:
        w, h = img.size
    if preset:
        return watermark_region_from_preset(w, h, position, margin, region_w, region_h)
    return clamp_crop_box(custom_box, w, h)


def resize_image(source: Path, dest: Path, options: ResizeOptions) -> Path:
    img = Image.open(source)
    orig_w, orig_h = img.size

    if options.max_side:
        scale = options.max_side / max(orig_w, orig_h)
        if scale >= 1.0:
            new_w, new_h = orig_w, orig_h
        else:
            new_w = max(1, int(orig_w * scale))
            new_h = max(1, int(orig_h * scale))
    elif options.width and options.height:
        if options.keep_aspect:
            img = ImageOps.contain(img, (options.width, options.height), Image.Resampling.LANCZOS)
            return _save_image(img, dest, options.quality)
        new_w, new_h = options.width, options.height
    elif options.width:
        ratio = options.width / orig_w
        new_w = options.width
        new_h = max(1, int(orig_h * ratio))
    elif options.height:
        ratio = options.height / orig_h
        new_h = options.height
        new_w = max(1, int(orig_w * ratio))
    else:
        raise ValueError("请指定宽度、高度或最大边长")

    resized = img.resize((new_w, new_h), Image.Resampling.LANCZOS)
    return _save_image(resized, dest, options.quality)


def rotate_image(source: Path, dest: Path, angle: float, expand: bool = True) -> Path:
    img = Image.open(source)
    rotated = img.rotate(angle, expand=expand, resample=Image.Resampling.BICUBIC)
    return _save_image(rotated, dest)


def flip_image(source: Path, dest: Path, mode: Literal["horizontal", "vertical"]) -> Path:
    img = Image.open(source)
    if mode == "horizontal":
        flipped = ImageOps.mirror(img)
    else:
        flipped = ImageOps.flip(img)
    return _save_image(flipped, dest)


def crop_image(source: Path, dest: Path, box: CropBox) -> Path:
    img = Image.open(source)
    w, h = img.size
    left = max(0, min(box.left, w - 1))
    top = max(0, min(box.top, h - 1))
    right = max(left + 1, min(box.right, w))
    bottom = max(top + 1, min(box.bottom, h))
    cropped = img.crop((left, top, right, bottom))
    return _save_image(cropped, dest)


def export_ocr_txt(text: str, dest: Path) -> Path:
    dest.parent.mkdir(parents=True, exist_ok=True)
    dest.write_text(text, encoding="utf-8")
    return dest


def export_ocr_docx(text: str, dest: Path, source_name: str = "") -> Path:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("未安装 python-docx，请运行: pip install python-docx") from exc

    doc = Document()
    if source_name:
        doc.add_heading(f"OCR 结果 — {source_name}", level=1)
    for paragraph in text.splitlines() or [""]:
        doc.add_paragraph(paragraph)
    dest.parent.mkdir(parents=True, exist_ok=True)
    doc.save(dest)
    return dest


def process_batch_watermark(
    sources: list[Path],
    output_dir: Path | None,
    options: WatermarkOptions,
    *,
    suffix: str = "_watermarked",
) -> list[Path]:
    results: list[Path] = []
    batch = len(sources) > 1
    for src in sources:
        dest = resolve_output_path(src, output_dir, batch_mode=batch, suffix=suffix)
        results.append(add_watermark(src, dest, options))
    return results


def process_batch_remove_watermark(
    sources: list[Path],
    output_dir: Path | None,
    options: RemoveWatermarkOptions,
    *,
    box_resolver: Callable[[Path], CropBox],
    suffix: str = "_unwatermarked",
) -> list[Path]:
    results: list[Path] = []
    batch = len(sources) > 1
    for src in sources:
        dest = resolve_output_path(src, output_dir, batch_mode=batch, suffix=suffix)
        per_image = RemoveWatermarkOptions(
            box=box_resolver(src),
            inpaint_radius=options.inpaint_radius,
            method=options.method,
            mask_expand=options.mask_expand,
        )
        results.append(remove_watermark(src, dest, per_image))
    return results


def process_batch_resize(
    sources: list[Path],
    output_dir: Path | None,
    options: ResizeOptions,
    *,
    suffix: str = "_resized",
) -> list[Path]:
    results: list[Path] = []
    batch = len(sources) > 1
    for src in sources:
        dest = resolve_output_path(src, output_dir, batch_mode=batch, suffix=suffix)
        results.append(resize_image(src, dest, options))
    return results


def convert_image(source: Path, dest: Path, options: ConvertOptions) -> Path:
    ext = options.target_ext.lower()
    if ext not in {".jpg", ".jpeg", ".png", ".bmp", ".gif", ".webp", ".tiff", ".tif"}:
        raise ValueError(f"不支持的目标格式: {ext}")

    img = Image.open(source)
    if getattr(img, "n_frames", 1) > 1 and ext == ".gif":
        img.seek(0)

    if ext in (".jpg", ".jpeg"):
        dest = dest.with_suffix(".jpg")
    elif ext == ".tiff":
        dest = dest.with_suffix(".tiff")
    else:
        dest = dest.with_suffix(ext)

    prepared = _prepare_for_format(img, dest.suffix, options.jpeg_bg)
    return _save_image(prepared, dest, options.quality)


def compress_image(source: Path, dest: Path, options: CompressOptions) -> CompressResult:
    original_bytes = source.stat().st_size
    img = Image.open(source)
    if getattr(img, "n_frames", 1) > 1:
        img.seek(0)

    if options.max_side:
        w, h = img.size
        scale = options.max_side / max(w, h)
        if scale < 1.0:
            new_w = max(1, int(w * scale))
            new_h = max(1, int(h * scale))
            img = img.resize((new_w, new_h), Image.Resampling.LANCZOS)

    if options.force_jpeg:
        dest = dest.with_suffix(".jpg")
    elif dest.suffix.lower() not in IMAGE_EXTENSIONS:
        dest = dest.with_suffix(source.suffix)

    prepared = _prepare_for_format(img, dest.suffix, options.jpeg_bg)
    saved = _save_image(prepared, dest, options.quality)
    output_bytes = saved.stat().st_size
    return CompressResult(path=saved, original_bytes=original_bytes, output_bytes=output_bytes)


def process_batch_convert(
    sources: list[Path],
    output_dir: Path | None,
    options: ConvertOptions,
) -> list[Path]:
    results: list[Path] = []
    batch = len(sources) > 1
    ext = options.target_ext if options.target_ext != ".jpeg" else ".jpg"
    for src in sources:
        if src.suffix.lower() == ext and not batch:
            stem = src.stem + "_converted"
        else:
            stem = src.stem
        if output_dir is None:
            dest = src.parent / f"{stem}{ext}"
        else:
            dest = output_dir / f"{stem}{ext}"
        results.append(convert_image(src, dest, options))
    return results


def process_batch_compress(
    sources: list[Path],
    output_dir: Path | None,
    options: CompressOptions,
    *,
    suffix: str = "_compressed",
) -> list[CompressResult]:
    results: list[CompressResult] = []
    batch = len(sources) > 1
    for src in sources:
        out_ext = ".jpg" if options.force_jpeg else src.suffix
        dest = resolve_output_path(src, output_dir, batch_mode=batch, suffix=suffix, new_ext=out_ext)
        results.append(compress_image(src, dest, options))
    return results
