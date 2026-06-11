#!/usr/bin/env python3
"""将识别文本导出为 TXT / Word。"""

from __future__ import annotations

from pathlib import Path


def export_txt(text: str, dest: Path) -> Path:
    dest.parent.mkdir(parents=True, exist_ok=True)
    dest.write_text(text, encoding="utf-8")
    return dest


def export_docx(text: str, dest: Path, *, title: str = "") -> Path:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("未安装 python-docx，请运行: pip install python-docx") from exc

    doc = Document()
    if title:
        doc.add_heading(title, level=1)
    for paragraph in text.splitlines() or [""]:
        doc.add_paragraph(paragraph)
    dest.parent.mkdir(parents=True, exist_ok=True)
    doc.save(dest)
    return dest
